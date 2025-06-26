import io
import sqlite3
from datetime import datetime
from pprint import pprint
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import qrcode
import streamlit as st
from PIL import Image
from scipy.integrate import quad
from scipy.optimize import differential_evolution
from scipy.stats import poisson
from docx import Document
from docx.shared import Inches

# DISPLAY THE RANDOM LOGO AT THE TOP WITH REDUCED SIZE
st.image("random_logo.png", width=200)

# INITIALIZE SESSION STATE TO STORE SELECTED TEMPLATES
if "selected_template" not in st.session_state:
    st.session_state.selected_template = "PREVENTIVE MAINTENANCE"
if "selected_inventory_part" not in st.session_state:
    st.session_state.selected_inventory_part = None
if "selected_technician" not in st.session_state:
    st.session_state.selected_technician = None
if "form_active" not in st.session_state:
    st.session_state.form_active = False

# DATABASE SETUP
def init_db():
    print("CONNECTING TO DATABASE ...")
    conn = sqlite3.connect("reliability_data.db")
    c = conn.cursor()
    
    # CREATE WORK_ORDERS TABLE
    c.execute('''CREATE TABLE IF NOT EXISTS work_orders (
                 id INTEGER PRIMARY KEY AUTOINCREMENT,
                 template_type TEXT,
                 asset_id TEXT,
                 description TEXT,
                 priority TEXT,
                 requested_date TEXT
                 )''')
    
    # CREATE INVENTORY TABLE
    c.execute('''CREATE TABLE IF NOT EXISTS inventory (
                 id INTEGER PRIMARY KEY AUTOINCREMENT,
                 part_id TEXT,
                 name_description TEXT,
                 location TEXT,
                 quantity_on_hand INTEGER,
                 min_level INTEGER,
                 max_level INTEGER,
                 last_restock_date TEXT,
                 supplier_info TEXT
                 )''')
    
    # CREATE TECHNICIANS TABLE
    c.execute('''CREATE TABLE IF NOT EXISTS technicians (
                 id INTEGER PRIMARY KEY AUTOINCREMENT,
                 name TEXT,
                 technician_id TEXT,
                 contact_details TEXT,
                 certifications TEXT,
                 skill_sets TEXT,
                 experience_level TEXT,
                 work_location TEXT,
                 shift_schedule TEXT
                 )''')
    
    conn.commit()
    conn.close()
    print("DATABASE INITIALIZED\n")

# INITIALIZE DATABASE
init_db()

# LOAD WORK ORDERS FROM DATABASE
def load_work_orders():
    conn = sqlite3.connect("reliability_data.db")
    c = conn.cursor()
    c.execute("SELECT * FROM work_orders")
    work_orders = []
    for row in c.fetchall():
        work_orders.append({
            "id": row[0],
            "TEMPLATE TYPE": row[1],
            "ASSET ID": row[2],
            "DESCRIPTION": row[3],
            "PRIORITY": row[4],
            "REQUESTED DATE": row[5]
        })
    conn.close()
    return work_orders

# LOAD INVENTORY FROM DATABASE
def load_inventory():
    conn = sqlite3.connect("reliability_data.db")
    c = conn.cursor()
    c.execute("SELECT * FROM inventory")
    inventory = []
    for row in c.fetchall():
        inventory.append({
            "id": row[0],
            "PART ID/SKU": row[1],
            "NAME AND DESCRIPTION": row[2],
            "LOCATION": row[3],
            "QUANTITY ON HAND": row[4],
            "MIN LEVEL": row[5],
            "MAX LEVEL": row[6],
            "LAST RESTOCK DATE": row[7],
            "SUPPLIER INFO": row[8]
        })
    conn.close()
    return inventory

# LOAD TECHNICIANS FROM DATABASE
def load_technicians():
    conn = sqlite3.connect("reliability_data.db")
    c = conn.cursor()
    c.execute("SELECT * FROM technicians")
    technicians = []
    for row in c.fetchall():
        technicians.append({
            "id": row[0],
            "NAME": row[1],
            "TECHNICIAN ID": row[2],
            "CONTACT DETAILS": row[3],
            "CERTIFICATIONS": row[4],
            "SKILL SETS": row[5],
            "EXPERIENCE LEVEL": row[6],
            "WORK LOCATION": row[7],
            "SHIFT SCHEDULE": row[8]
        })
    conn.close()
    return technicians

# FUNCTION TO GENERATE WORD DOCUMENT
def generate_word_document():
    work_orders = load_work_orders()
    inventory = load_inventory()
    technicians = load_technicians()
    
    doc = Document()
    doc.add_heading('MISSION RELIABILITY EVALUATOR - SAVED DATA', 0)
    
    # WORK ORDERS SECTION
    doc.add_heading('WORK ORDERS', level=1)
    if work_orders:
        for idx, order in enumerate(work_orders):
            doc.add_heading(f'WORK ORDER {idx + 1} - {order["TEMPLATE TYPE"]}', level=2)
            doc.add_paragraph(f"TEMPLATE TYPE: {order['TEMPLATE TYPE']}")
            doc.add_paragraph(f"ASSET ID: {order['ASSET ID']}")
            doc.add_paragraph(f"DESCRIPTION: {order['DESCRIPTION']}")
            doc.add_paragraph(f"PRIORITY: {order['PRIORITY']}")
            doc.add_paragraph(f"REQUESTED DATE: {order['REQUESTED DATE']}")
            doc.add_paragraph()  # ADD SPACING
    else:
        doc.add_paragraph("NO WORK ORDERS SAVED.")
    
    # INVENTORY SECTION
    doc.add_heading('INVENTORY', level=1)
    if inventory:
        for idx, item in enumerate(inventory):
            doc.add_heading(f'PART {idx + 1} - {item["PART ID/SKU"]}', level=2)
            doc.add_paragraph(f"PART ID/SKU: {item['PART ID/SKU']}")
            doc.add_paragraph(f"NAME AND DESCRIPTION: {item['NAME AND DESCRIPTION']}")
            doc.add_paragraph(f"LOCATION: {item['LOCATION']}")
            doc.add_paragraph(f"QUANTITY ON HAND: {item['QUANTITY ON HAND']}")
            doc.add_paragraph(f"MIN LEVEL: {item['MIN LEVEL']}")
            doc.add_paragraph(f"MAX LEVEL: {item['MAX LEVEL']}")
            doc.add_paragraph(f"LAST RESTOCK DATE: {item['LAST RESTOCK DATE']}")
            doc.add_paragraph(f"SUPPLIER INFO: {item['SUPPLIER INFO']}")
            if item['QUANTITY ON HAND'] < item['MIN LEVEL']:
                doc.add_paragraph("WARNING: QUANTITY ON HAND IS BELOW THE MINIMUM LEVEL!", style='Intense Quote')
            doc.add_paragraph()
    else:
        doc.add_paragraph("NO INVENTORY ITEMS SAVED.")
    
    # TECHNICIANS SECTION
    doc.add_heading('TECHNICIAN PROFILES', level=1)
    if technicians:
        for idx, tech in enumerate(technicians):
            doc.add_heading(f'TECHNICIAN {idx + 1} - {tech["NAME"]}', level=2)
            doc.add_paragraph(f"NAME: {tech['NAME']}")
            doc.add_paragraph(f"TECHNICIAN ID: {tech['TECHNICIAN ID']}")
            doc.add_paragraph(f"CONTACT DETAILS: {tech['CONTACT DETAILS']}")
            doc.add_paragraph(f"CERTIFICATIONS & LICENSES: {tech['CERTIFICATIONS']}")
            doc.add_paragraph(f"SKILL SETS: {tech['SKILL SETS']}")
            doc.add_paragraph(f"EXPERIENCE LEVEL: {tech['EXPERIENCE LEVEL']}")
            doc.add_paragraph(f"WORK LOCATION/ZONE: {tech['WORK LOCATION']}")
            doc.add_paragraph(f"SHIFT SCHEDULE AND AVAILABILITY: {tech['SHIFT SCHEDULE']}")
            doc.add_paragraph()
    else:
        doc.add_paragraph("NO TECHNICIAN PROFILES SAVED.")
    
    # SAVE DOCUMENT TO A BYTES BUFFER
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# STREAMLIT APP LAYOUT
st.title("MISSION RELIABILITY EVALUATOR")

# LIST OF FEATURES WITH ADDED SPECIAL FEATURES
features = {
    "WORK ORDER MANAGEMENT": "CREATE, TRACK, AND MANAGE WORK ORDERS EFFICIENTLY TO ENSURE TIMELY COMPLETION OF TASKS. SUPPORTS VARIOUS TASK TYPES, INCLUDING PREVENTIVE MAINTENANCE, REACTIVE MAINTENANCE, INSPECTIONS. FIELDS TYPICALLY INCLUDE: ASSET/EQUIPMENT ID, DESCRIPTION OF ISSUE/TASK, PRIORITY LEVEL, REQUESTED DATE/TIME.",
    "ASSET & EQUIPMENT MANAGEMENT": "MONITOR AND MANAGE ALL ASSETS AND EQUIPMENT, INCLUDING THEIR MAINTENANCE HISTORY AND PERFORMANCE.",
    "PREVENTIVE MAINTENANCE": "SCHEDULE AND AUTOMATE PREVENTIVE MAINTENANCE TASKS TO REDUCE UNEXPECTED BREAKDOWNS.",
    "VENDOR MANAGEMENT": "MANAGE VENDOR INFORMATION, CONTRACTS, AND COMMUNICATIONS TO STREAMLINE EXTERNAL SUPPORT.",
    "MAINTENANCE REQUESTS": "ALLOW STAFF TO SUBMIT MAINTENANCE REQUESTS EASILY AND TRACK THEIR STATUS IN REAL-TIME.",
    "PARTS & INVENTORY MANAGEMENT": "TRACK INVENTORY LEVELS, MANAGE SPARE PARTS, AND REORDER SUPPLIES TO AVOID DELAYS.",
    "RESOURCE & LABOR MANAGEMENT": "ASSIGN RESOURCES AND LABOR TO TASKS, ENSURING OPTIMAL WORKFORCE UTILIZATION.",
    "REPORTS & KPIS": "GENERATE DETAILED REPORTS AND KEY PERFORMANCE INDICATORS TO MONITOR SYSTEM PERFORMANCE.",
    "WEB & MOBILE APP": "ACCESS THE SYSTEM VIA WEB OR MOBILE APP FOR ON-THE-GO MANAGEMENT AND UPDATES.",
    "QR CODE SCANNING": "USE QR CODES TO QUICKLY ACCESS ASSET DETAILS, WORK ORDERS, OR MAINTENANCE HISTORY.",
    "LIFETIME CUSTOMER SUPPORT": "GET UNLIMITED SUPPORT FROM OUR TEAM TO ENSURE SMOOTH OPERATION OF YOUR SYSTEM.",
    "DOWNTIME TRACKING": "MONITOR AND ANALYZE EQUIPMENT DOWNTIME TO IDENTIFY PATTERNS AND REDUCE FUTURE OCCURRENCES.",
    "COST ANALYSIS": "TRACK MAINTENANCE COSTS AND ANALYZE EXPENSES TO OPTIMIZE BUDGET ALLOCATION.",
    "CUSTOM NOTIFICATIONS": "SET UP CUSTOM ALERTS FOR UPCOMING MAINTENANCE, OVERDUE TASKS, OR LOW INVENTORY LEVELS.",
    "PREDICTIVE MAINTENANCE": "LEVERAGE AI TO PREDICT EQUIPMENT FAILURES BEFORE THEY OCCUR, MINIMIZING DOWNTIME.",
    "AI-DRIVEN INSIGHTS": "GAIN ACTIONABLE INSIGHTS FROM DATA ANALYSIS TO OPTIMIZE MAINTENANCE STRATEGIES.",
    "AUTOMATED SCHEDULING": "AUTOMATICALLY SCHEDULE MAINTENANCE TASKS BASED ON EQUIPMENT USAGE AND PRIORITY.",
    "TECHNICIAN PROFILES & SKILLS TRACKING": "TRACK DETAILED INFO ABOUT EACH TECHNICIAN, INCLUDING SKILLS, CERTIFICATIONS, AND AVAILABILITY."
}

# WORK ORDER MANAGEMENT SECTION
if st.button("WORK ORDER MANAGEMENT") or st.session_state.get("form_active", False):
    st.session_state.form_active = True
    print("\n(BTN) WORK ORDER BUTTON CLICKED")
    st.write("**WORK ORDER MANAGEMENT**: CREATE, TRACK, AND MANAGE WORK ORDERS EFFICIENTLY TO ENSURE TIMELY COMPLETION OF TASKS. INCLUDES EDITABLE TEMPLATES FORM FOR PREVENTIVE MAINTENANCE, REACTIVE MAINTENANCE, INSPECTIONS. FIELDS TYPICALLY INCLUDE ASSET/EQUIPMENT ID, DESCRIPTION OF ISSUE/TASK, PRIORITY LEVEL, REQUESTED DATE/TIME.")
    st.subheader("MANAGE WORK ORDERS")

    # FORM FOR CREATING OR EDITING WORK ORDERS
    with st.form(key="work_order_form"):
        print("CREATING WORK ORDER FORM...")
        selected_template = st.selectbox("SELECT TEMPLATE TYPE", ["INSPECTIONS", "REACTIVE MAINTENANCE", "PREVENTIVE MAINTENANCE"], index=["INSPECTIONS", "REACTIVE MAINTENANCE", "PREVENTIVE MAINTENANCE"].index(st.session_state.selected_template), key="template_type")
        asset_id = st.text_input("ASSET/EQUIPMENT ID", value=st.session_state.get("asset_id", "EQ001"), key="asset_id_input")
        description = st.text_area("DESCRIPTION OF ISSUE/TASK", value=st.session_state.get("description", "ROUTINE CHECK-UP"), key="description_input")
        priority = st.selectbox("PRIORITY LEVEL", ["LOW", "MEDIUM", "HIGH"], index=["LOW", "MEDIUM", "HIGH"].index(st.session_state.get("priority", "MEDIUM")), key="priority_input")
        requested_date = st.text_input("REQUESTED DATE/TIME (YYYY-MM-DD HH:MM)", value=st.session_state.get("requested_date", "2025-06-04 09:00"), key="requested_date_input")
        submit_button = st.form_submit_button(label="SAVE WORK ORDER")

        if submit_button:
            print("(BTN) THE SUBMIT BUTTON WAS CLICKED")
            # VALIDATE DATE
            try:
                datetime.strptime(requested_date, "%Y-%m-%d %H:%M")
            except ValueError:
                st.error("PLEASE ENTER THE REQUESTED DATE/TIME IN YYYY-MM-DD HH:MM FORMAT.")
                st.stop()

            # CREATE OR UPDATE WORK ORDER
            work_order = {
                "TEMPLATE TYPE": selected_template,
                "ASSET ID": asset_id,
                "DESCRIPTION": description,
                "PRIORITY": priority,
                "REQUESTED DATE": requested_date
            }
            
            pprint(work_order)
            conn = sqlite3.connect("reliability_data.db")
            c = conn.cursor()
            if "edit_index" in st.session_state:
                c.execute('''UPDATE work_orders SET template_type = ?, asset_id = ?, description = ?, priority = ?, requested_date = ? WHERE id = ?''',
                          (selected_template, asset_id, description, priority, requested_date, st.session_state.edit_index))
                del st.session_state["edit_index"]
            else:
                c.execute('''INSERT INTO work_orders (template_type, asset_id, description, priority, requested_date) VALUES (?, ?, ?, ?, ?)''',
                          (selected_template, asset_id, description, priority, requested_date))
            conn.commit()
            conn.close()
            st.session_state.selected_template = selected_template
            st.session_state["asset_id"] = asset_id
            st.session_state["description"] = description
            st.session_state["priority"] = priority
            st.session_state["requested_date"] = requested_date
            st.success("WORK ORDER SAVED SUCCESSFULLY!")

    # DISPLAY AND EDIT/DELETE SAVED WORK ORDERS
    work_orders = load_work_orders()
    if work_orders:
        st.subheader("SAVED WORK ORDERS")
        for idx, order in enumerate(work_orders):
            with st.expander(f"WORK ORDER {idx + 1} - {order['TEMPLATE TYPE']}"):
                st.write(f"- **TEMPLATE TYPE**: {order['TEMPLATE TYPE']}")
                st.write(f"- **ASSET ID**: {order['ASSET ID']}")
                st.write(f"- **DESCRIPTION**: {order['DESCRIPTION']}")
                st.write(f"- **PRIORITY**: {order['PRIORITY']}")
                st.write(f"- **REQUESTED DATE**: {order['REQUESTED DATE']}")
                col1, col2 = st.columns(2)
                if col1.button("EDIT", key=f"edit_{idx}"):
                    st.session_state.edit_index = order["id"]
                    st.session_state["template_type"] = order["TEMPLATE TYPE"]
                    st.session_state["asset_id"] = order["ASSET ID"]
                    st.session_state["description"] = order["DESCRIPTION"]
                    st.session_state["priority"] = order["PRIORITY"]
                    st.session_state["requested_date"] = order["REQUESTED DATE"]
                    st.experimental_rerun()
                if col2.button("DELETE", key=f"delete_{idx}"):
                    conn = sqlite3.connect("reliability_data.db")
                    c = conn.cursor()
                    c.execute("DELETE FROM work_orders WHERE id = ?", (order["id"],))
                    conn.commit()
                    conn.close()
                    st.experimental_rerun()

# PARTS & INVENTORY MANAGEMENT SECTION
if st.button("PARTS & INVENTORY MANAGEMENT") or st.session_state.get("form_active", False):
    st.session_state.form_active = True
    print("\n(BTN) PARTS & INVENTORY MANAGEMENT BUTTON CLICKED")
    st.write("**PARTS & INVENTORY MANAGEMENT**: TRACK INVENTORY LEVELS, MANAGE SPARE PARTS, AND REORDER SUPPLIES TO AVOID DELAYS.")
    st.subheader("INVENTORY TRACKING")

    # FORM FOR ADDING OR EDITING INVENTORY ITEMS
    with st.form(key="inventory_form"):
        print("CREATING INVENTORY FORM...")
        part_id = st.text_input("PART ID/SKU", value=st.session_state.get("part_id", "PART001"), key="part_id_input")
        name_description = st.text_area("NAME AND DESCRIPTION", value=st.session_state.get("name_description", "VALVE PLUG/DISK - STANDARD SIZE"), key="name_description_input")
        location = st.text_input("LOCATION (WAREHOUSE, SITE, TRUCK, ETC.)", value=st.session_state.get("location", "WAREHOUSE A"), key="location_input")
        quantity_on_hand = st.number_input("QUANTITY ON HAND", min_value=0, value=st.session_state.get("quantity_on_hand", 10), key="quantity_on_hand_input")
        min_level = st.number_input("MINIMUM LEVEL", min_value=0, value=st.session_state.get("min_level", 5), key="min_level_input")
        max_level = st.number_input("MAXIMUM LEVEL", min_value=0, value=st.session_state.get("max_level", 20), key="max_level_input")
        last_restock_date = st.text_input("LAST RESTOCK DATE (YYYY-MM-DD)", value=st.session_state.get("last_restock_date", "2025-05-01"), key="last_restock_date_input")
        supplier_info = st.text_area("SUPPLIER/VENDOR INFORMATION", value=st.session_state.get("supplier_info", "SUPPLIER: ABC CORP\nCONTACT: 555-1234"), key="supplier_info_input")
        submit_button = st.form_submit_button(label="SAVE INVENTORY ITEM")

        if submit_button:
            print("(BTN) THE SUBMIT BUTTON WAS CLICKED")
            # VALIDATE INPUTS
            try:
                datetime.strptime(last_restock_date, "%Y-%m-%d")
            except ValueError:
                st.error("PLEASE ENTER THE LAST REST Suggested change: RESTOCK DATE IN YYYY-MM-DD FORMAT.")
                st.stop()

            # CREATE OR UPDATE INVENTORY ITEM
            inventory_item = {
                "PART ID/SKU": part_id,
                "NAME AND DESCRIPTION": name_description,
                "LOCATION": location,
                "QUANTITY ON HAND": quantity_on_hand,
                "MIN LEVEL": min_level,
                "MAX LEVEL": max_level,
                "LAST RESTOCK DATE": last_restock_date,
                "SUPPLIER INFO": supplier_info
            }
            pprint(inventory_item)
            conn = sqlite3.connect("reliability_data.db")
            c = conn.cursor()
            if "edit_inventory_index" in st.session_state:
                c.execute('''UPDATE inventory SET part_id = ?, name_description = ?, location = ?, quantity_on_hand = ?, min_level = ?, max_level = ?, last_restock_date = ?, supplier_info = ? WHERE id = ?''',
                          (part_id, name_description, location, quantity_on_hand, min_level, max_level, last_restock_date, supplier_info, st.session_state.edit_inventory_index))
                del st.session_state["edit_inventory_index"]
            else:
                c.execute('''INSERT INTO inventory (part_id, name_description, location, quantity_on_hand, min_level, max_level, last_restock_date, supplier_info) VALUES (?, ?, ?, ?, ?, ?, ?, ?)''',
                          (part_id, name_description, location, quantity_on_hand, min_level, max_level, last_restock_date, supplier_info))
            conn.commit()
            conn.close()
            st.session_state["part_id"] = part_id
            st.session_state["name_description"] = name_description
            st.session_state["location"] = location
            st.session_state["quantity_on_hand"] = quantity_on_hand
            st.session_state["min_level"] = min_level
            st.session_state["max_level"] = max_level
            st.session_state["last_restock_date"] = last_restock_date
            st.session_state["supplier_info"] = supplier_info
            st.success("INVENTORY ITEM SAVED SUCCESSFULLY!")

    # DISPLAY AND EDIT/DELETE INVENTORY ITEMS
    inventory = load_inventory()
    if inventory:
        st.subheader("CURRENT INVENTORY")
        for idx, item in enumerate(inventory):
            with st.expander(f"PART {idx + 1} - {item['PART ID/SKU']}"):
                st.write(f"- **PART ID/SKU**: {item['PART ID/SKU']}")
                st.write(f"- **NAME AND DESCRIPTION**: {item['NAME AND DESCRIPTION']}")
                st.write(f"- **LOCATION**: {item['LOCATION']}")
                st.write(f"- **QUANTITY ON HAND**: {item['QUANTITY ON HAND']}")
                st.write(f"- **MIN LEVEL**: {item['MIN LEVEL']}")
                st.write(f"- **MAX LEVEL**: {item['MAX LEVEL']}")
                st.write(f"- **LAST RESTOCK DATE**: {item['LAST RESTOCK DATE']}")
                st.write(f"- **SUPPLIER INFO**: {item['SUPPLIER INFO']}")
                if item['QUANTITY ON HAND'] < item['MIN LEVEL']:
                    st.warning("QUANTITY ON HAND IS BELOW THE MINIMUM LEVEL!")
                col1, col2 = st.columns(2)
                if col1.button("EDIT", key=f"edit_inventory_{idx}"):
                    st.session_state.edit_inventory_index = item["id"]
                    st.session_state["part_id"] = item["PART ID/SKU"]
                    st.session_state["name_description"] = item["NAME AND DESCRIPTION"]
                    st.session_state["location"] = item["LOCATION"]
                    st.session_state["quantity_on_hand"] = item["QUANTITY ON HAND"]
                    st.session_state["min_level"] = item["MIN LEVEL"]
                    st.session_state["max_level"] = item["MAX LEVEL"]
                    st.session_state["last_restock_date"] = item["LAST RESTOCK DATE"]
                    st.session_state["supplier_info"] = item["SUPPLIER INFO"]
                    st.experimental_rerun()
                if col2.button("DELETE", key=f"delete_inventory_{idx}"):
                    conn = sqlite3.connect("reliability_data.db")
                    c = conn.cursor()
                    c.execute("DELETE FROM inventory WHERE id = ?", (item["id"],))
                    conn.commit()
                    conn.close()
                    st.experimental_rerun()

# TECHNICIAN PROFILES & SKILLS TRACKING SECTION
if st.button("TECHNICIAN PROFILES & SKILLS TRACKING") or st.session_state.get("form_active", False):
    st.session_state.form_active = True
    print("\n(BTN) TECHNICIAN PROFILES & SKILLS TRACKING BUTTON CLICKED")
    st.write("**TECHNICIAN PROFILES & SKILLS TRACKING**: TRACK DETAILED INFO ABOUT EACH TECHNICIAN, INCLUDING SKILLS, CERTIFICATIONS, AND AVAILABILITY.")
    st.subheader("MANAGE TECHNICIANS")

    # FORM FOR ADDING OR EDITING TECHNICIANS
    with st.form(key="technician_form"):
        print("CREATING TECHNICIAN FORM...")
        name = st.text_input("NAME", value=st.session_state.get("technician_name", "JOHNSON DOE"), key="technician_name_input")
        technician_id = st.text_input("TECHNICIAN ID", value=st.session_state.get("technician_id", "TECH001"), key="technician_id_input")
        contact_details = st.text_area("CONTACT DETAILS", value=st.session_state.get("contact_details", "PHONE: 555-1234\nEMAIL: JOHNSON.DOE@EXAMPLE.COM"), key="contact_details_input")
        certifications = st.text_area("CERTIFICATIONS & LICENSES", value=st.session_state.get("certifications", "CERTIFIED MECHANICAL MAINTENANCE TECHNICIAN"), key="certifications_input")
        skill_sets = st.text_area("SKILL SETS (E.G., WELDING, ELECTRICAL, CORROSION ASSESSMENT)", value=st.session_state.get("skill_sets", "EQUIPMENT REPAIR AND INSTALLATION, SAFETY COMPLIANCE"), key="skill_sets_input")
        experience_level = st.selectbox("EXPERIENCE LEVEL", ["ENTRY LEVEL", "INTERMEDIATE", "SENIOR"], index=["ENTRY LEVEL", "INTERMEDIATE", "SENIOR"].index(st.session_state.get("experience_level", "INTERMEDIATE")), key="experience_level_input")
        work_location = st.text_input("WORK LOCATION/ZONE", value=st.session_state.get("work_location", "ZONE A"), key="work_location_input")
        shift_schedule = st.text_area("SHIFT SCHEDULE AND AVAILABILITY", value=st.session_state.get("shift_schedule", "MON-FRI, 8 AM - 4 PM\nAVAILABLE FOR OVERTIME"), key="shift_schedule_input")
        submit_button = st.form_submit_button(label="SAVE TECHNICIAN PROFILE")

        if submit_button:
            print("(BTN) THE SUBMIT BUTTON WAS CLICKED")
            # CREATE OR UPDATE TECHNICIAN PROFILE
            technician = {
                "NAME": name,
                "TECHNICIAN ID": technician_id,
                "CONTACT DETAILS": contact_details,
                "CERTIFICATIONS": certifications,
                "SKILL SETS": skill_sets,
                "EXPERIENCE LEVEL": experience_level,
                "WORK LOCATION": work_location,
                "SHIFT SCHEDULE": shift_schedule
            }
            pprint(technician)
            conn = sqlite3.connect("reliability_data.db")
            c = conn.cursor()
            if "edit_technician_index" in st.session_state:
                c.execute('''UPDATE technicians SET name = ?, technician_id = ?, contact_details = ?, certifications = ?, skill_sets = ?, experience_level = ?, work_location = ?, shift_schedule = ? WHERE id = ?''',
                          (name, technician_id, contact_details, certifications, skill_sets, experience_level, work_location, shift_schedule, st.session_state.edit_technician_index))
                del st.session_state["edit_technician_index"]
            else:
                c.execute('''INSERT INTO technicians (name, technician_id, contact_details, certifications, skill_sets, experience_level, work_location, shift_schedule) VALUES (?, ?, ?, ?, ?, ?, ?, ?)''',
                          (name, technician_id, contact_details, certifications, skill_sets, experience_level, work_location, shift_schedule))
            conn.commit()
            conn.close()
            st.session_state["technician_name"] = name
            st.session_state["technician_id"] = technician_id
            st.session_state["contact_details"] = contact_details
            st.session_state["certifications"] = certifications
            st.session_state["skill_sets"] = skill_sets
            st.session_state["experience_level"] = experience_level
            st.session_state["work_location"] = work_location
            st.session_state["shift_schedule"] = shift_schedule
            st.success("TECHNICIAN PROFILE SAVED SUCCESSFULLY!")

    # DISPLAY AND EDIT/DELETE TECHNICIAN PROFILES
    technicians = load_technicians()
    if technicians:
        st.subheader("TECHNICIAN PROFILES")
        for idx, tech in enumerate(technicians):
            with st.expander(f"TECHNICIAN {idx + 1} - {tech['NAME']}"):
                st.write(f"- **NAME**: {tech['NAME']}")
                st.write(f"- **TECHNICIAN ID**: {tech['TECHNICIAN ID']}")
                st.write(f"- **CONTACT DETAILS**: {tech['CONTACT DETAILS']}")
                st.write(f"- **CERTIFICATIONS & LICENSES**: {tech['CERTIFICATIONS']}")
                st.write(f"- **SKILL SETS**: {tech['SKILL SETS']}")
                st.write(f"- **EXPERIENCE LEVEL**: {tech['EXPERIENCE LEVEL']}")
                st.write(f"- **WORK LOCATION/ZONE**: {tech['WORK LOCATION']}")
                st.write(f"- **SHIFT SCHEDULE AND AVAILABILITY**: {tech['SHIFT SCHEDULE']}")
                col1, col2 = st.columns(2)
                if col1.button("EDIT", key=f"edit_technician_{idx}"):
                    st.session_state.edit_technician_index = tech["id"]
                    st.session_state["technician_name"] = tech["NAME"]
                    st.session_state["technician_id"] = tech["TECHNICIAN ID"]
                    st.session_state["contact_details"] = tech["CONTACT DETAILS"]
                    st.session_state["certifications"] = tech["CERTIFICATIONS"]
                    st.session_state["skill_sets"] = tech["SKILL SETS"]
                    st.session_state["experience_level"] = tech["EXPERIENCE LEVEL"]
                    st.session_state["work_location"] = tech["WORK LOCATION"]
                    st.session_state["shift_schedule"] = tech["SHIFT SCHEDULE"]
                    st.experimental_rerun()
                if col2.button("DELETE", key=f"delete_technician_{idx}"):
                    conn = sqlite3.connect("reliability_data.db")
                    c = conn.cursor()
                    c.execute("DELETE FROM technicians WHERE id = ?", (tech["id"],))
                    conn.commit()
                    conn.close()
                    st.experimental_rerun()

# DOWNLOAD SAVED DATA AS WORD DOCUMENT
if st.button("DOWNLOAD SAVED DATA AS WORD DOCUMENT"):
    st.header("DOWNLOAD ALL SAVED DATA")
    st.write("CLICK BELOW TO DOWNLOAD ALL SAVED WORK ORDERS, INVENTORY ITEMS, AND TECHNICIAN PROFILES AS A WORD DOCUMENT.")
    
    # GENERATE THE WORD DOCUMENT
    doc_buffer = generate_word_document()
    
    # PROVIDE DOWNLOAD BUTTON
    st.download_button(
        label="DOWNLOAD WORD DOCUMENT",
        data=doc_buffer,
        file_name="Reliability_Data.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

# QR CODE GENERATION SECTION FOR ALL SAVED DATA
if st.button("GENERATE QR CODE FOR ALL SAVED DATA"):
    st.header("QR CODE FOR ALL SAVED DATA")
    st.write("SCAN THIS QR CODE TO VIEW ALL SAVED WORK ORDERS, INVENTORY ITEMS, AND TECHNICIAN PROFILES.")

    # LOAD ALL DATA
    work_orders = load_work_orders()
    inventory = load_inventory()
    technicians = load_technicians()

    # FORMAT DATA INTO A CONCISE STRING
    data_lines = []
    for order in work_orders:
        data_lines.append(f"WO:{order['TEMPLATE TYPE']},{order['ASSET ID']},{order['PRIORITY']},{order['REQUESTED DATE']}")
    for item in inventory:
        data_lines.append(f"INV:{item['PART ID/SKU']},{item['NAME AND DESCRIPTION']},{item['QUANTITY ON HAND']},{item['LOCATION']}")
    for tech in technicians:
        data_lines.append(f"TECH:{tech['NAME']},{tech['TECHNICIAN ID']},{tech['SKILL SETS']},{tech['WORK LOCATION']}")
    qr_data = "\n".join(data_lines) if data_lines else "NO DATA SAVED."

    # GENERATE QR CODE
    qr = qrcode.QRCode(
        version=1,
        error_correction=qrcode.constants.ERROR_CORRECT_L,
        box_size=5,
        border=2,
    )
    qr.add_data(qr_data)
    qr.make(fit=True)

    # CREATE AN IMAGE FROM THE QR CODE INSTANCE
    img = qr.make_image(fill_color="black", back_color="white")

    # SAVE THE IMAGE TO A BYTES BUFFER TO DISPLAY IN STREAMLIT
    img_buffer = io.BytesIO()
    img.save(img_buffer, format="PNG")
    img_buffer.seek(0)

    # DISPLAY THE QR CODE IN STREAMLIT WITH A FIXED WIDTH
    st.image(img_buffer, caption="QR CODE CONTAINING ALL SAVED DATA", width=200)

# PARAMETERS INPUT
st.header("PARAMETERS")
col1, col2 = st.columns(2)

with col1:
    W = st.number_input("WORKLOAD (W, GOPS)", value=600.0)
    g = st.number_input("PROCESSING SPEED (G, GOPS/HR)", value=12.0)
    lambda_ = st.number_input("SHOCK RATE (LAMBDA, /HR)", value=0.25, step=0.01)
    lambda_tilde = st.number_input("SHOCK RATE DURING RESCUE (LAMBDA_TILDE, /HR)", value=0.06, step=0.01)
    alpha = st.number_input("SHARING FACTOR (ALPHA)", value=0.8, step=0.1)
    w = st.number_input("INSPECTION WORKLOAD (W, GOPS)", value=25.0)
    epsilon = st.number_input("DETECTION CUTOFF (EPSILON)", value=0.6, step=0.1)
    p = st.number_input("FALSE NEGATIVE RATE (P)", value=0.05, step=0.01)
    S_star_85 = st.number_input("CONSTRAINT S >= 0.85 (S*)", value=0.85, step=0.01)

with col2:
    q = st.number_input("FALSE POSITIVE RATE (Q)", value=0.03, step=0.01)
    delta = st.number_input("DECELERATION DURING RESCUE (DELTA)", value=0.7, step=0.1)
    mu0 = st.number_input("BASE RESCUE TIME (MU_0, HR)", value=12.0)
    mu1 = st.number_input("RESCUE TIME GROWTH (MU_1)", value=0.1, step=0.1)
    eta = st.number_input("WEIBULL SCALE (ETA, HR)", value=120.0)
    beta = st.number_input("WEIBULL SHAPE (BETA)", value=2.0)
    max_m = st.number_input("MAX SHOCKS (M_MAX)", value=10, step=1, format="%d")
    N = st.number_input("NUMBER OF INSPECTIONS (N)", value=1, step=1, format="%d")
    S_star_90 = st.number_input("CONSTRAINT S >= 0.90 (S*)", value=0.90, step=0.01)

# FUNCTIONS
def theta():
    return w / (alpha * g)

def total_mission_time(N):
    return (W + N * w) / g

def phi_i(tau, i):
    return mu0 + mu1 * (tau[i-1] * g + w / alpha - i * w) / W

def z(k):
    return 1 if k == 0 else 0.97 * (0.85) ** (k - 1)

def Z(m):
    return np.prod([z(i) for i in range(m + 1)])

def P(t, m, lambda_val):
    return poisson.pmf(m, lambda_val * t)

def u(t):
    return lambda_ * sum(P(t, m-1, lambda_) * (1 - z(m)) * Z(m-1) for m in range(1, max_m))

def u_tilde(t, tau_i, theta_val):
    return lambda_tilde * sum(
        P(tau_i + theta_val, k, lambda_) * 
        sum(P(t, l, lambda_tilde) * Z(k + l) for l in range(max_m))
        for k in range(max_m)
    )

def V(t):
    return 1 - np.exp(-((t / eta) ** beta)) if t > 0 else 0

def calculate_mission_success_probability(tau, N, T, theta_val):
    tau = [0] + (list(tau) if np.isscalar(tau) or len(tau) > 0 else []) + [T]
    R = 0
    if N > 0:
        R += (1 - q) ** N * sum(P(T, m, lambda_) * Z(m) for m in range(max_m))
        for i in range(1, N + 1):
            integral, _ = quad(lambda t: (1 - V(T - t)) * u(t), 
                              tau[i-1] + epsilon * theta_val, 
                              tau[i] + epsilon * theta_val)
            R += (1 - q) ** (i-1) * p ** (N - i + 1) * integral
        integral, _ = quad(lambda t: (1 - V(T - t)) * u(t), 
                          tau[N] + epsilon * theta_val, T)
        R += (1 - q) ** N * integral
    else:
        R += sum(P(T, m, lambda_) * Z(m) for m in range(max_m))
        integral, _ = quad(lambda t: (1 - V(T - t)) * u(t), 0, T)
        R += integral
    return R

def calculate_failure_avoidance_probability(tau, N, T, theta_val):
    tau = [0] + (list(tau) if np.isscalar(tau) or len(tau) > 0 else []) + [T]
    S = 0
    for i in range(1, N + 1):
        phi = phi_i(tau, i)
        for k in range(1, i + 1):
            integral, _ = quad(
                lambda t: (1 - V(tau[i] + theta_val - t + delta * phi)) * u(t),
                tau[k-1] + epsilon * theta_val, 
                tau[k] + epsilon * theta_val)
            S += (1 - q) ** (k-1) * p ** (i-k) * (1 - p) * integral
    for i in range(1, N + 1):
        phi = phi_i(tau, i)
        term1 = sum(P(tau[i] + theta_val, k, lambda_) * 
                    sum(P(phi, l, lambda_tilde) * Z(k + l) for l in range(max_m))
                    for k in range(max_m))
        integral, _ = quad(
            lambda t: (1 - V(delta * (phi - t))) * u_tilde(t, tau[i], theta_val),
            0, phi)
        S += q * (1 - q) ** (i-1) * (term1 + integral)
    S += calculate_mission_success_probability(tau[1:-1], N, T, theta_val)
    return S

def objective_1(lambda_val):
    global lambda_
    lambda_ = lambda_val
    N = 0
    theta_val = theta()
    T = total_mission_time(N)
    R = calculate_mission_success_probability([], N, T, theta_val)
    S = calculate_failure_avoidance_probability([], N, T, theta_val)
    return R, S

def objective_2(N, lambda_val):
    global lambda_
    lambda_ = lambda_val
    theta_val = theta()
    T = total_mission_time(N)
    def objective_de(tau):
        return -calculate_failure_avoidance_probability([tau[0]], N, T, theta_val)
    result_de = differential_evolution(objective_de, bounds=[(0, T)], maxiter=50)
    tau_de = [result_de.x[0]]
    S_de = -result_de.fun
    R_de = calculate_mission_success_probability(tau_de, N, T, theta_val)
    return R_de, S_de, tau_de

def objective_3(lambda_val):
    global lambda_
    lambda_ = lambda_val
    N = 1
    theta_val = theta()
    T = total_mission_time(N)
    def objective_de(tau):
        S = calculate_failure_avoidance_probability([tau[0]], N, T, theta_val)
        R = calculate_mission_success_probability([tau[0]], N, T, theta_val)
        penalty = 1e6 * max(0, S_star_90 - S)
        return -R + penalty
    result_de = differential_evolution(objective_de, bounds=[(0, T)], maxiter=50)
    tau_de = [result_de.x[0]]
    R_de = calculate_mission_success_probability(tau_de, N, T, theta_val)
    S_de = calculate_failure_avoidance_probability(tau_de, N, T, theta_val)
    return R_de, S_de, tau_de

def objective_4(lambda_val):
    global lambda_
    lambda_ = lambda_val
    N = 1
    theta_val = theta()
    T = total_mission_time(N)
    def objective_de(tau):
        S = calculate_failure_avoidance_probability([tau[0]], N, T, theta_val)
        R = calculate_mission_success_probability([tau[0]], N, T, theta_val)
        penalty = 1e6 * max(0, S_star_85 - S)
        return -R + penalty
    result_de = differential_evolution(objective_de, bounds=[(0, T)], maxiter=50)
    tau_de = [result_de.x[0]]
    R_de = calculate_mission_success_probability(tau_de, N, T, theta_val)
    S_de = calculate_failure_avoidance_probability(tau_de, N, T, theta_val)
    return R_de, S_de, tau_de

# CALCULATE BUTTON
if st.button("CALCULATE"):
    st.header("RESULTS")
    
    # OBJECTIVE 1
    st.subheader("OBJECTIVE 1: R = S, NO INSPECTIONS")
    R, S = objective_1(lambda_)
    st.write(f"MISSION SUCCESS PROBABILITY (R): {R:.3f}")
    st.write(f"FAILURE AVOIDANCE PROBABILITY (S): {S:.3f}")
    
    # OBJECTIVE 2
    st.subheader("OBJECTIVE 2: MAXIMIZE S")
    R_de, S_de, tau_de = objective_2(N, lambda_)
    st.write(f"MISSION SUCCESS PROBABILITY (R): {R_de:.3f}")
    st.write(f"FAILURE AVOIDANCE PROBABILITY (S): {S_de:.3f}")
    st.write(f"OPTIMAL INSPECTION TIME (TAU_1): {tau_de[0]:.3f} HR")
     
   
    # OBJECTIVE 3
    st.subheader("OBJECTIVE 3: MAXIMIZE R S.T. S >= 0.90")
    R_de, S_de, tau_de = objective_3(lambda_)
    st.write(f"MISSION SUCCESS PROBABILITY (R): {R_de:.3f}")
    st.write(f"FAILURE AVOIDANCE PROBABILITY (S): {S_de:.3f}")
    st.write(f"OPTIMAL INSPECTION TIME (TAU_1): {tau_de[0]:.3f} HR")
    
    # OBJECTIVE 4
    st.subheader("OBJECTIVE 4: MAXIMIZE R S.T. S >= 0.85")
    R_de, S_de, tau_de = objective_4(lambda_)
    st.write(f"MISSION SUCCESS PROBABILITY (R): {R_de:.3f}")
    st.write(f"FAILURE AVOIDANCE PROBABILITY (S): {S_de:.3f}")
    st.write(f"OPTIMAL INSPECTION TIME (TAU_1): {tau_de[0]:.3f} HR")

# ADD DATE AND TIME AND COPYRIGHT NOTICE AT THE BOTTOM
st.write("LAST UPDATED: WEDNESDAY, JUNE 25, 2025, 06:57 AM -03")
st.write("© 2025 ALL RIGHTS RESERVED.")
